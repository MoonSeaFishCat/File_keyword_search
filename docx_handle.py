import os
from docx import Document
import pandas as pd
from collections import defaultdict
import re


def normalize_school_name(school):
    # 这个函数可以根据需要进行修改，目的是生成一个更加通用的匹配模式
    # 例如，可以将所有空白转换为\s+，这样可以匹配任何数量的空白字符
    return re.sub(r'\s+', r'\\s+', school)


def find_schools_in_docx(docx_path, schools):
    doc = Document(docx_path)
    school_count = defaultdict(int)

    # 组合文档中的所有文本
    text = "\n".join([para.text for para in doc.paragraphs])
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text += "\n" + cell.text

    text = text.lower()  # 将文本转换为小写，匹配时不区分大小写

    # 对于每一个学校，生成一个模糊匹配的正则表达式并计数
    for school in schools:
        normalized_school = normalize_school_name(school)
        pattern = re.compile(normalized_school, flags=re.I)
        school_count[school] += len(pattern.findall(text))

    return school_count


current_directory = os.getcwd()
data_file_path = os.path.join(current_directory, 'data.txt')

with open(data_file_path, 'r', encoding='utf-8') as file:
    schools = [line.strip() for line in file.readlines()]

results = []
for root, dirs, files in os.walk(os.path.join(current_directory, 'input')):
    print(root)
    for file in files:
        if file.lower().endswith('.docx'):
            docx_path = os.path.join(root, file)
            relative_folder = os.path.relpath(root, start=current_directory)
            school_count = find_schools_in_docx(docx_path, schools)
            for school, count in school_count.items():
                if count > 0:
                    results.append({
                        'File Path': docx_path,
                        'Folder': relative_folder,
                        'School': school,
                        'Count': count
                    })

df = pd.DataFrame(results)
output_directory = os.path.join(current_directory, 'output')
if not os.path.exists(output_directory):
    os.makedirs(output_directory)

excel_path = os.path.join(output_directory, 'schools_count_docx.xlsx')
df.to_excel(excel_path, index=False)
